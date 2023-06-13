from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from hydra import initialize, compose
from docx import *
from docx.shared import Cm
from collections import Counter
from functions.book.book_operations import count_words


def create_report(book,cfg):

    book.chapters.remove(book.chapters[45])
    book.chapters.remove(book.chapters[44])
    book.chapters.remove(book.chapters[43])
    book.chapters.remove(book.chapters[0])
    book.chapters.remove(book.chapters[0])
    book.chapters.remove(book.chapters[0])
    book.chapters.remove(book.chapters[0])
    book.chapters.remove(book.chapters[0])
    book.chapters.remove(book.chapters[66])

    doc = Document()

    # Creating the title page
    doc.add_heading(book.title, 0)
    doc.add_picture(cfg.picture.path_edited, width=Cm(10))
    doc.add_paragraph(book.author)
    doc.add_paragraph(f"Report by {cfg.params.author}")

    # Create pages with plots
    for idx, chapter in enumerate(book.chapters, 1):
        words_count = count_words(chapter)

        title = chapter.select('h4')
        title = str(title)
        title = title.replace("[<h4>", "")
        title = title.replace("</h4>]", "")

        plt.figure()
        counter = Counter(words_count)
        x_values = list(counter.keys())
        y_values = list(counter.values())

        plt.bar(x_values, y_values)
        plt.yticks(np.arange(min(y_values), max(y_values)+1,1.0))
        plt.xlabel('Word Counts (rounded to nearest 10)')
        plt.ylabel('Number of Paragraphs')
        plt.title(f'Distribution of Paragraph Lengths in Chapter {idx}')
        plt.savefig('temp.png')
        doc.add_page_break()
        doc.add_paragraph(f"Chapter {idx}: {title}")
        doc.add_picture('temp.png')
        os.remove('temp.png')
        plt.close()

    # Creating final page
    doc.add_page_break()
    num_paragraphs = [len(i) for i in book.chapters]
    plt.plot(num_paragraphs)
    plt.xlim([1,70])
    plt.xlabel('Chapter')
    plt.ylabel('Number of Paragraphs')
    plt.title('Number of Paragraphs in Each Chapter')
    plt.savefig('temp.png')
    doc.add_picture('temp.png', width=Cm(15))
    os.remove('temp.png')
    plt.close()

    # Add metrics table
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chapter'
    hdr_cells[1].text = 'Min Words'
    hdr_cells[2].text = 'Max Words'
    hdr_cells[3].text = 'Average Words'

    for idx, chapter in enumerate(book.chapters, 1):
        words_count = count_words(chapter)
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = str(np.min(words_count))
        cells[2].text = str(np.max(words_count))
        cells[3].text = str(round(np.mean(words_count),2))

    # Save the document with date and book id in the filename
    doc.save(f"Report_{book.id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")